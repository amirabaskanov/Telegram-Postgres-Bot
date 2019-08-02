from lib2to3.pgen2.grammar import line

import telebot
from telebot import types
import psycopg2
from docx import Document
from docx.shared import Pt, RGBColor
bot = telebot.TeleBot('939294233:AAF32a0UlQztHM8IULCM9nAmRmkytWw6c5Y')

document = Document('Resume.docx')

first_name = ''
last_name = ''
age = 0


@bot.message_handler(content_types=['text'])
def start(message):
    if message.text == '/hi' or message.text == '/start':
        # postgres_create_query = 'CREATE TABLE IF NOT EXISTS cv (chat_id INT PRIMARY KEY , first_name VARCHAR(40), last_name VARCHAR(40), age VARCHAR(40);'

        a = str(message.from_user)
        bot.send_message(message.from_user.id, "What's Your name?")
        bot.register_next_step_handler(message, get_first_name)
        with open('cv.doc', 'a') as f:
            f.write(a + '\n')
    else:
        bot.send_message(message.from_user.id, 'Write /hi')


def get_first_name(message):
    global first_name
    a = first_name = message.text
    bot.send_message(message.from_user.id, "What's Your surname?")
    bot.register_next_step_handler(message, get_last_name)
    with open('cv.doc', 'a') as f:
        f.write(a + '\n')


def get_last_name(message):
    global last_name
    a = last_name = message.text
    bot.send_message(message.from_user.id, "How old are you?")
    bot.register_next_step_handler(message, get_age)
    with open('cv.doc', 'a') as f:
        f.write(a + '\n')


def get_age(message):
    global age
    if age == 0:
        try:
            age = int(message.text)
            keyboard = types.InlineKeyboardMarkup()
            key_yes = types.InlineKeyboardButton(text='Yes', callback_data='yes')
            keyboard.add(key_yes)
            key_no = types.InlineKeyboardButton(text='No', callback_data='no')
            keyboard.add(key_no)
            question = 'You are ' + str(age) + ' years old, Your name is ' + first_name + \
                       ' ' + last_name + '.' + ' Want to save it?'
            bot.send_message(message.from_user.id, text=question, reply_markup=keyboard)

        except Exception:
            bot.send_message(message.from_user.id, 'use numbers, please)')

        with open('cv.doc', 'a') as f:
            f.write(str(age) + '\n')

    else:
        pass


@bot.callback_query_handler(func=lambda call: True)
def callback_worker(call):
    if call.data == 'yes':
        bot.send_message(call.message.chat.id, 'Remember it: )')
        try:
            connection = psycopg2.connect(user="amir",
                                          password="password1",
                                          host="localhost",
                                          port="5432",
                                          database="cv")
            cursor = connection.cursor()
            # items = []
            # print(items)
            # items.append(call.message.chat.id)
            # items.append(first_name)
            # items.append(last_name)
            # items.append(age)
            postgres_insert_query = 'INSERT INTO cv ("chat_id", "first_name", "last_name", "age") VALUES (%s,%s,%s,%s)'
            record_to_insert = (call.message.chat.id, first_name, last_name, age)
            cursor.execute(postgres_insert_query, record_to_insert, )

            cursor = connection.cursor()
            cursor.execute("SELECT * FROM cv where chat_id = '%s';" % call.message.chat.id)
            results = cursor.fetchall()
            connection.commit()

            for row in results:


                # print("first_name = ", row[0], )
                # print("last_name = ", row[1], )
                # print("age  = ", row[2], )
                # print("chat_id  = ", row[3], "\n")

                first_name_doc = document.tables[0].rows[0].cells[0].paragraphs[0]
                first_name_doc.style.font.size = Pt(24)
                first_name_doc.style.font.bold = True

                last_name_doc = document.tables[0].rows[0].cells[0].paragraphs[1]
                last_name_doc.style.font.size = Pt(24)
                last_name_doc.style.font.bold = True

                age_doc = document.tables[0].rows[0].cells[0].paragraphs[2]
                age_doc.style.font.size = Pt(12)
                age_doc.style.font.bold = False

                first_name_doc.text = row[0]
                last_name_doc.text = row[1]
                age_doc.text = row[2]

                document.save('Resume_edited.docx')

        except (Exception, psycopg2.Error) as error:
            if (connection):
                print("Failed to insert record into mobile table", error)

    elif call.data == 'no':
        bot.send_message(call.message.chat.id, 'Delete it: )')



bot.polling(none_stop=True, interval=0)
