from docx import Document
from docx.shared import Pt, RGBColor

document = Document('Resume.docx')



first_name = document.tables[0].rows[0].cells[0].paragraphs[0]
first_name.style.font.size = Pt(24)
first_name.style.font.bold = True


last_name = document.tables[0].rows[0].cells[0].paragraphs[1]
last_name.style.font.size = Pt(24)
last_name.style.font.bold = True

age = document.tables[0].rows[0].cells[0].paragraphs[2]
age.style.font.size = Pt(12)
age.style.font.bold = False




first_name.text = Dictionary["first_name"]
last_name.text = Dictionary["last_name"]
age.text = Dictionary["age"]


document.save('Resume_edited.docx')